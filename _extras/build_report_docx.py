import re
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


ROOT = Path(__file__).resolve().parent
MARKDOWN_PATH = ROOT / "report_text.md"
OUTPUT_PATH = ROOT / "ECON3350_Research_Report.docx"

FIGURE_FALLBACKS = {
    "Figure 1": ["fig1_log_levels.png"],
    "Figure 2": ["fig2_log_diffs.png"],
    "Figure 3": ["fig2a_forecast.png"],
    "Figures 3a–3c": [
        "fig2b_arima_3_0_3.png",
        "fig2b_arima_1_0_6.png",
        "fig2b_arima_5_0_6.png",
    ],
    "Figures 3a-3c": [
        "fig2b_arima_3_0_3.png",
        "fig2b_arima_1_0_6.png",
        "fig2b_arima_5_0_6.png",
    ],
    "Figures 4-6": [
        "fig3_actual_vs_arima_3_0_3.png",
        "fig3_actual_vs_arima_1_0_6.png",
        "fig3_actual_vs_arima_5_0_6.png",
    ],
    "Figures 4–6": [
        "fig3_actual_vs_arima_3_0_3.png",
        "fig3_actual_vs_arima_1_0_6.png",
        "fig3_actual_vs_arima_5_0_6.png",
    ],
    "Figure 7": ["fig4a_real_rate.png"],
    "Figure 8": ["fig4b_consumption_ratio.png"],
    "Figure 9": ["fig5b_abs_returns.png"],
    "Figure 10": ["fig6_sqstd_acf_pval.png"],
    "Figures 11-14": [
        "fig6_vol_CNY.png",
        "fig6_vol_USD.png",
        "fig6_vol_TWI.png",
        "fig6_vol_SDR.png",
    ],
    "Figures 11–14": [
        "fig6_vol_CNY.png",
        "fig6_vol_USD.png",
        "fig6_vol_TWI.png",
        "fig6_vol_SDR.png",
    ],
}


def set_default_style(document: Document) -> None:
    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3", "Title"):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(12)
        style.font.color.rgb = RGBColor(0, 0, 0)
        if style_name != "Normal":
            style.font.bold = True


def set_run_style(run, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def math_segments(text: str) -> list[tuple[str, bool, bool]]:
    """
    Parse a math expression into (segment_text, is_subscript, is_superscript).

    Handles LaTeX-style notation:
      x_t        → x normal, t subscript
      x_{T+1}    → x normal, T+1 subscript
      x^2        → x normal, 2 superscript
      x^{abc}    → x normal, abc superscript

    Unicode sub/superscript characters (₁ ² etc.) are left as-is in normal runs.
    """
    segments: list[tuple[str, bool, bool]] = []
    i = 0
    base = ""
    n = len(text)

    while i < n:
        c = text[i]

        if c == '_' and i + 1 < n:
            if base:
                segments.append((base, False, False))
                base = ""
            i += 1
            if i < n and text[i] == '{':
                try:
                    end = text.index('}', i + 1)
                    segments.append((text[i + 1:end], True, False))
                    i = end + 1
                except ValueError:
                    base += '_'
            else:
                # Single-character subscript
                segments.append((text[i], True, False))
                i += 1

        elif c == '^' and i + 1 < n:
            if base:
                segments.append((base, False, False))
                base = ""
            i += 1
            if i < n and text[i] == '{':
                try:
                    end = text.index('}', i + 1)
                    segments.append((text[i + 1:end], False, True))
                    i = end + 1
                except ValueError:
                    base += '^'
            else:
                segments.append((text[i], False, True))
                i += 1

        else:
            base += c
            i += 1

    if base:
        segments.append((base, False, False))
    return segments


SUBSCRIPT_MAP = str.maketrans({
    "0": "₀", "1": "₁", "2": "₂", "3": "₃", "4": "₄",
    "5": "₅", "6": "₆", "7": "₇", "8": "₈", "9": "₉",
    "+": "₊", "-": "₋", "=": "₌", "(": "₍", ")": "₎",
    "a": "ₐ", "e": "ₑ", "h": "ₕ", "i": "ᵢ", "j": "ⱼ",
    "k": "ₖ", "l": "ₗ", "m": "ₘ", "n": "ₙ", "o": "ₒ",
    "p": "ₚ", "r": "ᵣ", "s": "ₛ", "t": "ₜ", "u": "ᵤ",
    "v": "ᵥ", "x": "ₓ",
    ",": ",", "T": "T", "C": "C", "Y": "Y",
})

SUPERSCRIPT_MAP = str.maketrans({
    "0": "⁰", "1": "¹", "2": "²", "3": "³", "4": "⁴",
    "5": "⁵", "6": "⁶", "7": "⁷", "8": "⁸", "9": "⁹",
    "+": "⁺", "-": "⁻", "=": "⁼", "(": "⁽", ")": "⁾",
    "i": "ⁱ", "n": "ⁿ", "j": "ʲ",
})


def linear_math_text(text: str) -> str:
    """
    Convert simple LaTeX-ish sub/superscript markers to visible Unicode
    characters before placing the result inside a Word OMML equation object.
    This keeps formulas readable in Word without relying on code-font runs.
    """
    parts = math_segments(text)
    converted: list[str] = []
    for seg_text, is_sub, is_sup in parts:
        if is_sub:
            converted.append(seg_text.translate(SUBSCRIPT_MAP))
        elif is_sup:
            converted.append(seg_text.translate(SUPERSCRIPT_MAP))
        else:
            converted.append(seg_text)
    return "".join(converted)


def is_math_like(text: str) -> bool:
    math_markers = (
        "_", "^", "Δ", "δ", "μ", "σ", "ω", "α", "β", "φ", "θ",
        "Φ", "∈", "−", "×", "<", ">", "=", "sum ", "P(",
    )
    if any(marker in text for marker in math_markers):
        return True
    return bool(re.match(r"^[a-zA-Z]+\([^)]+\)(-[A-Z]+\([^)]+\))?$", text))


SCRIPTED_TOKEN = re.compile(
    r"([A-Za-z]+|[αβγδμσφθωΣΦΔ])"
    r"(?:(?:_\{([^}]+)\})|_([A-Za-z0-9,+\-−]+))?"
    r"(?:(?:\^\{([^}]+)\})|\^([A-Za-z0-9,+\-−]+))?"
)


def math_run(text: str):
    m_run = OxmlElement("m:r")
    m_text = OxmlElement("m:t")
    m_text.text = text
    m_run.append(m_text)
    return m_run


def scripted_math_element(base: str, subscript: str | None, superscript: str | None):
    if subscript and superscript:
        elem = OxmlElement("m:sSubSup")
        base_tag = "m:e"
        sub_tag = "m:sub"
        sup_tag = "m:sup"
    elif subscript:
        elem = OxmlElement("m:sSub")
        base_tag = "m:e"
        sub_tag = "m:sub"
        sup_tag = None
    elif superscript:
        elem = OxmlElement("m:sSup")
        base_tag = "m:e"
        sub_tag = None
        sup_tag = "m:sup"
    else:
        return math_run(base)

    base_elem = OxmlElement(base_tag)
    base_elem.append(math_run(base))
    elem.append(base_elem)

    if subscript:
        sub_elem = OxmlElement(sub_tag)
        sub_elem.append(math_run(subscript))
        elem.append(sub_elem)

    if superscript:
        sup_elem = OxmlElement(sup_tag)
        sup_elem.append(math_run(superscript))
        elem.append(sup_elem)

    return elem


def add_omml_math(paragraph, raw_text: str) -> None:
    o_math = OxmlElement("m:oMath")
    position = 0
    for match in SCRIPTED_TOKEN.finditer(raw_text):
        base = match.group(1)
        subscript = match.group(2) or match.group(3)
        superscript = match.group(4) or match.group(5)
        if match.start() > position:
            o_math.append(math_run(raw_text[position:match.start()]))
        o_math.append(scripted_math_element(base, subscript, superscript))
        position = match.end()
    if position < len(raw_text):
        o_math.append(math_run(raw_text[position:]))
    paragraph._p.append(o_math)


PLAIN_MATH_TOKEN = re.compile(r"(R²|[δμσ][̂]?(?:²)?|[αβφθωΦΣ][̂]?(?:[₀-₉ᵢⱼₜ₊₋]+)?|[A-Za-z]?[₀-₉ₜⱼᵢ₊₋]+(?:²)?)")


def add_plain_runs(paragraph, text: str) -> None:
    for piece in PLAIN_MATH_TOKEN.split(text):
        if not piece:
            continue
        if PLAIN_MATH_TOKEN.fullmatch(piece):
            add_omml_math(paragraph, piece)
        else:
            run = paragraph.add_run(piece)
            set_run_style(run)


def add_runs(paragraph, raw_text: str) -> None:
    """
    Parse raw_text and add appropriately formatted runs to paragraph.

    Recognised markers:
      **text**   → bold run
      `expr`     → Word equation object when mathematical; otherwise italic run
      plain text → normal run (inherits document style)
    """
    # Split into: bold spans, math spans, plain text
    pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`)')
    parts = pattern.split(raw_text)

    for part in parts:
        if not part:
            continue

        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            set_run_style(run, bold=True)

        elif part.startswith('`') and part.endswith('`'):
            inner = part[1:-1]
            if is_math_like(inner):
                add_omml_math(paragraph, inner)
            else:
                run = paragraph.add_run(inner)
                set_run_style(run, italic=True)

        else:
            add_plain_runs(paragraph, part)


def format_cell(cell, raw_text: str, bold_header: bool = False) -> None:
    """
    Replace the content of a table cell with math-aware formatted runs.
    Clears any existing runs before writing.
    """
    para = cell.paragraphs[0]
    # Remove all existing w:r (run) elements from the paragraph XML
    for elem in list(para._p.findall(qn('w:r'))):
        para._p.remove(elem)
    add_runs(para, raw_text)
    if bold_header:
        for run in para.runs:
            run.bold = True


def add_title_page(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("ECON3350 Research Report: Michael Nguyen (s4884533)")
    set_run_style(run, bold=True)


def add_markdown_table(document: Document, rows: list[str]) -> None:
    def split_row(row: str) -> list[str]:
        return [cell.strip() for cell in row.strip().strip("|").split("|")]

    data = [split_row(r) for r in rows]
    if len(data) < 2:
        return
    header = data[0]
    body = data[2:]  # row index 1 is the markdown separator (---|---:| etc.)

    table = document.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"

    # Header row — bold and math-aware
    hdr_cells = table.rows[0].cells
    for i, value in enumerate(header):
        format_cell(hdr_cells[i], value, bold_header=True)

    # Body rows
    for row in body:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if i < len(cells):
                format_cell(cells[i], value)

    document.add_paragraph("")


def add_figure_block(document: Document, line: str) -> None:
    # Extract PNG filenames from backtick-wrapped refs or bare refs
    image_names = re.findall(r"`([A-Za-z0-9_]+\.png)`", line)
    if not image_names:
        image_names = re.findall(r"([A-Za-z0-9_]+\.png)", line)
    if not image_names:
        stripped = re.sub(r"\*\*", "", line).strip().strip(":").strip()
        image_names = FIGURE_FALLBACKS.get(stripped, [])

    if image_names:
        for image_name in image_names:
            image_path = ROOT.parent / image_name
            if image_path.exists():
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(str(image_path), width=Inches(6.4))
        # Build caption: strip PNG refs, markdown bold/code markers
        caption = re.sub(r"`[A-Za-z0-9_]+\.png`", "", line)
        caption = re.sub(r"[A-Za-z0-9_]+\.png", "", caption)
        caption = re.sub(r"\*\*", "", caption)
        caption = re.sub(r"`", "", caption)
        caption = caption.strip().strip(":").strip()
        if caption:
            cap_p = document.add_paragraph()
            add_runs(cap_p, caption)
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = document.add_paragraph()
        add_runs(p, line)


def is_displayed_formula(line: str) -> bool:
    """
    Returns True if the entire (stripped) line is a single backtick-delimited
    expression — i.e. a standalone displayed formula, not inline math.
    """
    s = line.strip()
    return (
        len(s) > 2
        and s.startswith('`')
        and s.endswith('`')
        and s.count('`') == 2
    )


def build_docx() -> None:
    document = Document()
    set_default_style(document)
    add_title_page(document)

    lines = MARKDOWN_PATH.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Blank line
        if not line:
            i += 1
            continue

        # Page break
        if line == "---":
            document.add_page_break()
            i += 1
            continue

        # Headings
        if line.startswith("### "):
            document.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        if line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
            i += 1
            continue

        if line.startswith("# "):
            # Top-level title is on the title page; skip in body
            i += 1
            continue

        # Table block
        if line.startswith("|"):
            block = [line]
            j = i + 1
            while j < len(lines) and lines[j].startswith("|"):
                block.append(lines[j].rstrip())
                j += 1
            add_markdown_table(document, block)
            i = j
            continue

        # Figure block (line starts with **Figure or **[INSERT ...png...])
        if line.startswith("**Figure") or (line.startswith("**[INSERT") and ".png" in line):
            add_figure_block(document, line)
            i += 1
            continue

        # Displayed formula: entire line is one backtick-wrapped expression
        if is_displayed_formula(line):
            inner = line.strip()[1:-1]
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_omml_math(p, inner)
            i += 1
            continue

        # Numbered list item
        if re.match(r"^\d+\.", line):
            p = document.add_paragraph()
            add_runs(p, line)
            i += 1
            continue

        # Bullet list item
        if line.startswith("- "):
            p = document.add_paragraph(style="List Bullet")
            add_runs(p, line[2:])
            i += 1
            continue

        # Default: regular paragraph
        p = document.add_paragraph()
        add_runs(p, line)
        i += 1

    # Set document metadata
    now = datetime.now(timezone.utc)
    cp = document.core_properties
    cp.author = "Michael Nguyen"
    cp.last_modified_by = "Michael Nguyen"
    cp.created = now
    cp.modified = now

    document.save(str(OUTPUT_PATH))


if __name__ == "__main__":
    build_docx()
    print(f"Saved: {OUTPUT_PATH}")
